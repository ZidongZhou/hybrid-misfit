from __future__ import annotations

import argparse
import re
import shutil
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
MAN = ROOT / "manuscript"
TABLES = ROOT / "tables"
FIGS = ROOT / "figures"
SOURCE = ROOT / "documentation" / "current_sources" / "Manuscript_MDPI_Sustainability_source.md"
OFFICIAL_TEMPLATE = ROOT / "template" / "sustainability-template.docx"
OFFICIAL_DOT_TEMPLATE = ROOT / "template" / "sustainability-template.dot"

STYLE = {
    "article": "MDPI_1.1_article_type",
    "title": "MDPI_1.2_title",
    "authors": "MDPI_1.3_authornames",
    "affiliation": "MDPI_1.6_affiliation",
    "abstract": "MDPI_1.7_abstract",
    "keywords": "MDPI_1.8_keywords",
    "h1": "MDPI_2.1_heading1",
    "h2": "MDPI_2.2_heading2",
    "h3": "MDPI_2.3_heading3",
    "text": "MDPI_3.1_text",
    "text_no_indent": "MDPI_3.2_text_no_indent",
    "table_caption": "MDPI_4.1_table_caption",
    "table_body": "MDPI_4.2_table_body",
    "table_footer": "MDPI_4.3_table_footer",
    "figure": "MDPI_5.2_figure",
    "figure_caption": "MDPI_5.1_figure_caption",
    "back": "MDPI_6.2_back_matter",
    "notes": "MDPI_6.3_notes",
    "reference": "MDPI_8.1_references",
}

COL_LABELS = {
    "step": "Step", "n_unweighted": "N", "events": "Events", "weighted_prevalence": "Weighted prevalence",
    "unweighted_prevalence": "Unweighted prevalence", "misfit_group": "Misfit group", "n": "N",
    "job_dissatisfaction_weighted": "Dissatisfied", "female_share_weighted": "Female share",
    "children_share_weighted": "Children share", "income_mean_weighted": "Mean income",
    "commute_mean_weighted": "Mean commute, minutes", "education_mean_weighted": "Mean education category",
    "desired_wfh_days_weighted": "Desired WFH days", "planned_wfh_days_weighted": "Planned WFH days",
    "current_wfh_days_weighted": "Current WFH days", "feature_set": "Feature set", "model": "Model",
    "roc_auc": "ROC-AUC", "pr_auc": "PR-AUC", "brier": "Brier", "precision_at_10": "Precision@10%",
    "recall_at_10": "Recall@10%", "lift_at_10": "Lift@10%", "rule": "Rule", "method": "Method",
}

VALUE_LABELS = {
    "fit_or_weak": "Aligned or gap <2 days",
    "over_remote_strong": "Strong over-remote misfit",
    "under_remote_strong": "Strong under-remote misfit",
    "hris": "Conventional characteristics",
    "desired": "Conventional characteristics + desired WFH",
    "arrangement": "Conventional characteristics + desired/planned WFH",
    "restricted": "Broader planned-only feature set",
    "gender_children_excluded": "Reduced-feature",
    "gender_children_income_excluded": "Gender-, children-, and income-excluded",
    "rf": "Random forest",
    "logit": "Logit",
    "absolute_gap_rule": "Absolute-gap rule",
    "under_remote_directional_rule": "Under-remote directional rule",
    "additive_support_rule": "Additive rule",
    "gender_children_excluded_rf": "Reduced-feature random forest",
    "full_restricted_rf_comparison": "Broader random-forest comparison",
}

SUPPLEMENTARY_FILES = [
    ("Table S1. Sample construction flow", "revision_sample_construction_flow.csv", "S01_SampleFlow"),
    ("Table S2. Weighted descriptive statistics by misfit group", "revision_table1_descriptives_by_misfit.csv", "S02_Descriptives"),
    ("Table S3. Full weighted/unweighted association models", "revision_weighted_unweighted_association.csv", "S03_AssocBasic"),
    ("Table S4. Association model grid with controls", "revision_association_model_grid.csv", "S04_AssocControls"),
    ("Table S5. Single-imputation and missingness-indicator association robustness", "revision_imputed_association_robustness.csv", "S05_ImputedAssoc"),
    ("Table S6. Threshold and reference-arrangement robustness", "revision_threshold_reference_lpm_robustness.csv", "S06_ThresholdRef"),
    ("Table S7. Reference construction summary", "revision_reference_construction_summary.csv", "S07_Reference"),
    ("Table S8. Outcome-definition and misfit-definition robustness", "revision_outcome_misfit_definition_robustness.csv", "S08_OutcomeMisfit"),
    ("Table S9. Five-category satisfaction-score robustness", "revision_satisfaction_score_robustness.csv", "S09_SatisfactionScore"),
    ("Table S10. Planned-WFH missingness profile", "revision_planned_wfh_missingness_profile.csv", "S10_Missingness"),
    ("Table S11. Weight-trimming robustness", "revision_weight_trimming_lpm.csv", "S11_WeightTrim"),
    ("Table S12. Month-cluster robust uncertainty for the M4 model", "revision_wave_clustered_lpm.csv", "S12_ClusterSE"),
    ("Table S13. Temporal-holdout feature ablation", "revision_ablation_temporal_performance.csv", "S13_Validation"),
    ("Table S14. Directly interpretable rule definitions", "revision_rule_definitions.csv", "S14_Rules"),
    ("Table S15. Paired-bootstrap temporal-holdout comparisons", "revision_paired_bootstrap_vs_rules.csv", "S15_PairedBoot"),
    ("Table S16. Calibration correction metrics", "revision_calibration_correction.csv", "S16_Calibration"),
    ("Table S17. Top-k subgroup audit", "revision_topk_subgroup_audit.csv", "S17_SubgroupAudit"),
    ("Table S18. Weighted top-k sensitivity", "revision_weighted_topk_sensitivity.csv", "S18_WeightedTopK"),
    ("Table S19. Heterogeneity interaction results", "revision_heterogeneity_interactions.csv", "S19_Heterogeneity"),
    ("Table S20. Reverse-causality sensitivity checks", "revision_reverse_causality_sensitivity.csv", "S20_Reverse"),
    ("Table S21. Supplementary dissatisfaction-related outcomes", "revision_supplementary_outcomes_summary.csv", "S21_ProxyOutcomes"),
    ("Table S22. Outcome coding audit", "revision_outcome_coding_audit.csv", "S22_CodingAudit"),
    ("Table S23. Model tuning and preprocessing details", "revision_model_tuning_details.csv", "S23_Tuning"),
    ("Table S24. Practical ranking simulation", "revision_practical_ranking_simulation.csv", "S24_RankSim"),
    ("Table S25. Permutation importance", "revision_permutation_importance.csv", "S25_Importance"),
    ("Table S26. Continuous directional-misfit robustness", "revision_continuous_misfit_lpm.csv", "S26_Continuous"),
    ("Table S27. Complete-case inverse-probability weighting sensitivity", "revision_ipw_complete_case_sensitivity.csv", "S27_IPW"),
    ("Table S28. January-February 2026 temporal-holdout check", "revision_temporal_stability_validation.csv", "S28_Stability"),
    ("Table S29. Raw-score group-wise calibration diagnostics", "revision_groupwise_calibration.csv", "S29_GroupCal"),
    ("Table S30. Planned-only feature ablation", "revision_sensitive_variable_ablation.csv", "S30_FeatureAblation"),
    ("Table S31. Tie-robust top-k sensitivity", "revision_tie_robust_topk.csv", "S31_TieRobust"),
    ("Table S32. Calibrated subgroup audit", "revision_groupwise_calibrated_subgroup_audit.csv", "S32_CalFair"),
    ("Table S33. Remoteability moderation and stratification", "revision_remoteability_moderation.csv", "S33_Remoteability"),
    ("Table S34. Signed misfit-bin prevalence", "revision_misfit_bin_prevalence.csv", "S34_MisfitBins"),
    ("Table S35. Complete-case balance", "revision_complete_case_balance.csv", "S35_CCBalance"),
    ("Table S36. Planned-WFH missingness model", "revision_planned_missingness_model.csv", "S36_PlannedMiss"),
    ("Table S37. SWAA weight distribution by analytic sample", "revision_weight_distribution.csv", "S37_Weights"),
    ("Table S38. Planned-only ranking validation", "revision_planned_only_ranking_validation.csv", "S38_PlannedRank"),
    ("Table S39. Month-by-month January-February 2026 validation", "revision_monthly_2026q1_validation.csv", "S39_JanFeb"),
    ("Table S40. Weighted-training sensitivity", "revision_weighted_training_sensitivity.csv", "S40_WtdTraining"),
    ("Table S41. Ranking-model governance summary", "revision_ranking_model_governance_summary.csv", "S41_ModelGovernance"),
    ("Table S42. Subgroup safeguard simulation", "revision_subgroup_safeguard_simulation.csv", "S42_Safeguard"),
    ("Table S43. Planned-only polynomial response-surface analysis", "revision_response_surface_planned_only.csv", "S43_ResponseSurface"),
    ("Additional Table A1. Temporal calibration sensitivity", "revision_temporal_calibration_sensitivity.csv", "A1_TemporalCal"),
    ("Additional Table A2. Calibration-bin source data", "revision_calibration_bins.csv", "A2_CalibrationBins"),
    ("Additional Table A3. Primary association sample by survey month", "revision_association_sample_months.csv", "A3_AssocMonths"),
]


def ensure_template() -> Path:
    if OFFICIAL_TEMPLATE.exists():
        return OFFICIAL_TEMPLATE
    if not OFFICIAL_DOT_TEMPLATE.exists():
        raise FileNotFoundError("MDPI template not found")
    OFFICIAL_TEMPLATE.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(OFFICIAL_DOT_TEMPLATE, "r") as zin, zipfile.ZipFile(OFFICIAL_TEMPLATE, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(
                    b"application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                    b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                )
            zout.writestr(item, data)
    return OFFICIAL_TEMPLATE


def clear_body(doc: Document) -> Document:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    return doc


def new_mdpi_document() -> Document:
    return clear_body(Document(str(ensure_template())))


def has_style(doc: Document, style: str) -> bool:
    return style in {s.name for s in doc.styles}


def clean_text(text: str) -> str:
    """Remove Markdown-only delimiters before writing publication text to Word."""
    return str(text).replace("`", "")


def add_p(doc: Document, text: str = "", style: str | None = None, bold: bool = False, italic: bool = False,
          center: bool = False, keep_with_next: bool = False, keep_together: bool = False):
    text = clean_text(text)
    p = doc.add_paragraph(style=style if style and has_style(doc, style) else None)
    p.paragraph_format.keep_with_next = keep_with_next
    p.paragraph_format.keep_together = keep_together
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_labeled_p(doc: Document, label: str, text: str, style: str):
    label = clean_text(label)
    text = clean_text(text)
    p = doc.add_paragraph(style=style if has_style(doc, style) else None)
    r = p.add_run(label)
    r.bold = True
    if text:
        p.add_run(text)
    return p


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    tr_pr.append(el)


def no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.style = STYLE["table_body"] if has_style(p.part.document, STYLE["table_body"]) else p.style
    r = p.add_run(str(text))
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_value(col: str, val) -> str:
    if pd.isna(val):
        return ""
    if isinstance(val, (int,)) and not isinstance(val, bool):
        return f"{val:,}".replace("-", "−")
    if isinstance(val, float):
        if col in {"N", "Events", "Month clusters"}:
            return f"{val:,.0f}".replace("-", "−")
        if col in {"Weighted prevalence", "Unweighted prevalence", "Weighted dissatisfaction", "Female share", "Children share", "Complete-case retention"}:
            return f"{100 * val:.1f}%".replace("-", "−")
        if col in {"Mean income", "Mean commute, minutes", "Mean education category"}:
            return f"{val:.1f}".replace("-", "−")
        if col in {"Desired WFH days", "Planned WFH days", "Current WFH days"}:
            return f"{val:.2f}".replace("-", "−")
        if col in {"ROC-AUC", "PR-AUC", "Brier", "Precision@10%", "Recall@10%"}:
            return f"{val:.4f}".replace("-", "−")
        return f"{val:.3f}".replace("-", "−")
    return VALUE_LABELS.get(str(val), str(val))


def add_table(doc: Document, caption: str, df: pd.DataFrame, note: str | None = None, max_rows: int | None = None, page_break=False):
    if page_break:
        doc.add_page_break()
    add_p(doc, caption, STYLE["table_caption"], keep_with_next=True)
    d = df.copy().dropna(axis=1, how="all")
    if max_rows:
        d = d.head(max_rows)
    d = d.rename(columns={c: COL_LABELS.get(c, c) for c in d.columns})
    table = doc.add_table(rows=1, cols=len(d.columns))
    table.style = "MDPI_4.1_three_line_table" if "MDPI_4.1_three_line_table" in {s.name for s in doc.styles} else "Table Grid"
    repeat_header(table.rows[0])
    for j, col in enumerate(d.columns):
        set_cell_text(table.rows[0].cells[j], col, bold=True)
    for _, row in d.iterrows():
        tr = table.add_row()
        no_split(tr)
        for j, col in enumerate(d.columns):
            set_cell_text(tr.cells[j], format_value(str(col), row[col]))
    if note:
        add_labeled_p(doc, "Note: ", note, STYLE["table_footer"])
    return table


def p_value_text(x: float) -> str:
    if pd.isna(x):
        return ""
    return "<0.001" if x < 0.001 else f"{x:.3f}"


def est_ci(row: pd.Series) -> str:
    return f"{row['estimate']:.3f} [{row['ci_low']:.3f}, {row['ci_high']:.3f}]".replace("-", "−")


def add_table_1(doc, sample):
    labels = {
        "All rows in uploaded SWAA file": ("SWAA 2025–2026 rows", "Loaded registered-access SWAA release"),
        "Rows with job satisfaction outcome": ("Observed dissatisfaction", "Outcome observed"),
        "Rows with desired WFH days": ("Observed desired WFH", "Preference observed"),
        "Rows with planned-only misfit and ranking outcome": ("Observed desired/planned WFH and outcome", "Ranking-eligible observations"),
        "Complete-case association sample, September-December 2025": ("Complete-case association, Sep–Dec 2025", "Fully adjusted association model"),
        "2025 train analytic sample": ("Training sample, Jul–Dec 2025", "Training sample"),
        "January-February 2026 test analytic sample": ("Temporal holdout, Jan–Feb 2026", "Temporal holdout"),
    }
    d = sample[["step", "n_unweighted", "events", "weighted_prevalence"]].copy()
    d["Use"] = d["step"].map(lambda x: labels.get(x, (x, ""))[1])
    d["step"] = d["step"].map(lambda x: labels.get(x, (x, ""))[0])
    add_table(doc, "Table 1. Sample construction and analytic samples.", d,
              "Events and prevalence are shown when job dissatisfaction is observed. SWAA = Survey of Working Arrangements and Attitudes; WFH = work from home.")


def add_table_2(doc, desc):
    add_p(doc, "Table 2. Weighted descriptive statistics by planned-reference misfit group, July 2025–February 2026.", STYLE["table_caption"], keep_with_next=True)
    a = desc[["misfit_group", "n", "events", "job_dissatisfaction_weighted", "desired_wfh_days_weighted", "planned_wfh_days_weighted", "current_wfh_days_weighted"]]
    add_p(doc, "Panel A. Outcome and work-arrangement profile", STYLE["table_body"], bold=True, keep_with_next=True)
    add_table_body_only(doc, a)
    b = desc[["misfit_group", "female_share_weighted", "children_share_weighted", "income_mean_weighted", "commute_mean_weighted", "education_mean_weighted"]]
    add_p(doc, "Panel B. Worker characteristics", STYLE["table_body"], bold=True, keep_with_next=True)
    add_table_body_only(doc, b)
    add_labeled_p(doc, "Note: ", "WFH = work from home. Proportions and means are SWAA-weighted. Income is in USD thousands; commute is in minutes; education uses the SWAA numeric scale.", STYLE["table_footer"])


def add_table_body_only(doc, df):
    d = df.rename(columns={c: COL_LABELS.get(c, c) for c in df.columns})
    table = doc.add_table(rows=1, cols=len(d.columns))
    table.style = "MDPI_4.1_three_line_table" if "MDPI_4.1_three_line_table" in {s.name for s in doc.styles} else "Table Grid"
    repeat_header(table.rows[0])
    for j, col in enumerate(d.columns):
        set_cell_text(table.rows[0].cells[j], col, True)
    for _, row in d.iterrows():
        tr = table.add_row(); no_split(tr)
        for j, col in enumerate(d.columns):
            set_cell_text(tr.cells[j], format_value(str(col), row[col]))
    return table


def add_table_3(doc, assoc):
    rows = []
    for spec in ["M1 unadjusted", "M2 demographics family commute", "M3 conventional characteristics", "M4 conventional characteristics plus month FE"]:
        sub = assoc[(assoc["specification"] == spec) & (assoc["model"] == "weighted_lpm")]
        under = sub[sub["term"] == "under_remote_planned"].iloc[0]
        over = sub[sub["term"] == "over_remote_planned"].iloc[0]
        dif = sub[sub["term"] == "under_minus_over"].iloc[0]
        label = {
            "M1 unadjusted": "M1: Unadjusted",
            "M2 demographics family commute": "M2: Demographics + commute",
            "M3 conventional characteristics": "M3: Conventional characteristics",
            "M4 conventional characteristics plus month FE": "M4: Conventional + month FE",
        }[spec]
        rows.append({
            "Specification": label,
            "Under [95% CI]": est_ci(under),
            "Over [95% CI]": est_ci(over),
            "Difference [95% CI]": est_ci(dif),
            "p": p_value_text(float(dif["p_value"])),
            "N": int(under["n"]),
        })
    add_table(doc, "Table 3. Weighted linear probability model association summary.", pd.DataFrame(rows),
              "Confidence intervals are heteroskedasticity-robust. Month-cluster intervals are reported in Supplementary Table S12 because the main outcome window contains four monthly clusters.",
              page_break=True)


def add_table_4(doc, perf, rules):
    def fitted(feature, model):
        return perf[(perf["feature_set"] == feature) & (perf["model"] == model)].iloc[0]
    direction = rules[rules["rule"] == "under_remote_directional_rule"].iloc[0]
    rows = []
    for label, row in [
        ("Conventional logit", fitted("hris", "logit")),
        ("Desired/planned-WFH logit", fitted("arrangement", "logit")),
        ("Conventional RF", fitted("hris", "rf")),
        ("Desired/planned-WFH RF", fitted("arrangement", "rf")),
        ("Under-remote rule", direction),
        ("Reduced-feature RF", fitted("gender_children_excluded", "rf")),
    ]:
        rows.append({"Model": label, "ROC-AUC": row["roc_auc"], "PR-AUC": row["pr_auc"], "Precision@10%": row["precision_at_10"], "Recall@10%": row["recall_at_10"], "Lift@10%": row["lift_at_10"]})
    add_table(doc, "Table 4. January–February 2026 temporal-holdout performance.", pd.DataFrame(rows),
              "All fitted models use planned-only features. The reduced-feature RF excludes gender and children but retains income; current WFH is excluded from every fitted model. Supplementary Tables S13, S15, S30, and S41 report broader comparisons and paired-bootstrap uncertainty.",
              page_break=True)


def add_figure(doc, caption: str, filename: str, width=5.75):
    image = FIGS / filename
    if not image.exists():
        raise FileNotFoundError(image)
    p = doc.add_paragraph(style=STYLE["figure"] if has_style(doc, STYLE["figure"]) else None)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    run = p.add_run()
    shape = run.add_picture(str(image), width=Inches(width))
    shape._inline.docPr.set("title", caption.split(".")[0])
    shape._inline.docPr.set("descr", caption)
    cap = add_p(doc, caption, STYLE["figure_caption"], keep_together=True)
    return cap


def rendered_blocks():
    sample = pd.read_csv(TABLES / "revision_sample_construction_flow.csv")
    assoc = pd.read_csv(TABLES / "revision_association_model_grid.csv")
    perf = pd.read_csv(TABLES / "revision_ablation_temporal_performance.csv")
    rules = pd.read_csv(TABLES / "revision_strong_rule_baselines.csv")
    paired = pd.read_csv(TABLES / "revision_paired_bootstrap_vs_rules.csv")
    cal = pd.read_csv(TABLES / "revision_calibration_correction.csv")
    subgroup = pd.read_csv(TABLES / "revision_topk_subgroup_audit.csv")
    rsa = pd.read_csv(TABLES / "revision_response_surface_planned_only.csv")
    sr = lambda label: sample.loc[sample["step"] == label].iloc[0]
    eligible = sr("Nonmissing employer-planned WFH")
    complete, train, hold = sr("Complete-case association sample, September-December 2025"), sr("2025 train analytic sample"), sr("January-February 2026 test analytic sample")
    retention = complete["n_unweighted"] / eligible["n_unweighted"]
    sample_text=(f"This complete-case sample retained {100*retention:.1f}% of the {int(eligible['n_unweighted']):,} observations with job dissatisfaction, desired WFH, and employer-planned WFH observed. "
                 f"The ranking analysis used {int(train['n_unweighted']):,} observations from July–December 2025 for training and {int(hold['n_unweighted']):,} observations from January–February 2026 for testing; the holdout contained {int(hold['events']):,} dissatisfied respondents. Supplementary Table S1 and Additional Table A3 report the sample flow and month distribution.")
    m4=assoc[(assoc["specification"]=="M4 conventional characteristics plus month FE") & (assoc["model"]=="weighted_lpm")]
    u=m4[m4.term=="under_remote_planned"].iloc[0]; o=m4[m4.term=="over_remote_planned"].iloc[0]; c=m4[m4.term=="under_minus_over"].iloc[0]
    p_text=p_value_text(float(c["p_value"])); p_clause=f"p {p_text}" if p_text.startswith("<") else f"p = {p_text}"
    association=(f"Compared with respondents who were aligned or had a desired–planned gap of less than two days, strong under-remote misfit was associated with an {100*u['estimate']:.1f}-percentage-point higher probability of job dissatisfaction (95% CI: {100*u['ci_low']:.1f}–{100*u['ci_high']:.1f}). "
                 f"The corresponding over-remote estimate was {100*o['estimate']:.1f} percentage points (95% CI: {100*o['ci_low']:.1f}–{100*o['ci_high']:.1f}). The formal under–over contrast was {100*c['estimate']:.1f} percentage points (95% CI: {100*c['ci_low']:.1f}–{100*c['ci_high']:.1f}; {p_clause}).")
    def fitted(feature, model): return perf[(perf.feature_set==feature)&(perf.model==model)].iloc[0]
    hrf=fitted("hris","rf"); arf=fitted("arrangement","rf"); hlog=fitted("hris","logit"); alog=fitted("arrangement","logit"); red=fitted("gender_children_excluded","rf")
    rule=rules[rules.rule=="under_remote_directional_rule"].iloc[0]
    inc=paired[paired.comparison=="arrangement_rf_minus_conventional_rf"].iloc[0]
    inclog=paired[paired.comparison=="arrangement_logit_minus_conventional_logit"].iloc[0]
    cmp=paired[paired.comparison=="gender_children_excluded_rf_minus_under_remote_directional_rule"].iloc[0]
    abstract_rank=(f"Adding desired and planned WFH increased random-forest precision@10% from {hrf.precision_at_10:.4f} to {arf.precision_at_10:.4f}; the paired-bootstrap increase was {100*inc.delta_precision10_point:.2f} percentage points (95% interval: {100*inc.delta_precision10_low:.2f}–{100*inc.delta_precision10_high:.2f}). "
                   f"At 10% capacity, the under-remote rule achieved a precision of {rule.precision_at_10:.4f}, compared with {red.precision_at_10:.4f} for the reduced-feature random forest; the {100*cmp.delta_precision10_point:.2f}-percentage-point difference had a 95% interval of {str(f'{100*cmp.delta_precision10_low:.2f}').replace('-','−')} to {100*cmp.delta_precision10_high:.2f} percentage points.")
    model=(f"Adding desired and planned WFH increased random-forest precision@10% from {hrf.precision_at_10:.4f} to {arf.precision_at_10:.4f}. The paired-bootstrap increase was {100*inc.delta_precision10_point:.2f} percentage points (95% interval: {100*inc.delta_precision10_low:.2f}–{100*inc.delta_precision10_high:.2f}); PR-AUC increased by {inc.delta_pr_auc_point:.4f} (95% interval: {inc.delta_pr_auc_low:.4f}–{inc.delta_pr_auc_high:.4f}). "
           f"The corresponding logit comparison increased precision@10% from {hlog.precision_at_10:.4f} to {alog.precision_at_10:.4f}, with a paired-bootstrap increase of {100*inclog.delta_precision10_point:.2f} percentage points (95% interval: {100*inclog.delta_precision10_low:.2f}–{100*inclog.delta_precision10_high:.2f}). "
           f"The reduced-feature RF achieved ROC-AUC {red.roc_auc:.4f}, PR-AUC {red.pr_auc:.4f}, precision@10% {red.precision_at_10:.4f}, recall@10% {red.recall_at_10:.4f}, and lift@10% {red.lift_at_10:.3f}. The under-remote rule achieved precision@10% {rule.precision_at_10:.4f}; the RF–rule difference was {100*cmp.delta_precision10_point:.2f} percentage points (95% interval: {str(f'{100*cmp.delta_precision10_low:.2f}').replace('-','−')} to {100*cmp.delta_precision10_high:.2f} percentage points).")
    cm={r.model:r for _,r in cal.iterrows()}; raw=cm["raw_gender_children_excluded_rf"]; pl=cm["platt_gender_children_excluded_rf"]; iso=cm["isotonic_gender_children_excluded_rf"]; null=cm["null_2025_base_rate"]
    calibration=(f"The raw reduced-feature RF was poorly calibrated: its Brier score was {raw.brier:.4f}, compared with {null.brier:.4f} for the base-rate benchmark, and its ECE was {raw.ece:.4f}. Raw scores should therefore not be interpreted as probabilities. "
                 f"Platt and isotonic calibration reduced the Brier score to {pl.brier:.4f} and {iso.brier:.4f}, respectively, only {null.brier-pl.brier:.4f} and {null.brier-iso.brier:.4f} below the benchmark. These modest improvements do not establish a deployment-ready probability model; the primary comparison is therefore based on ranking metrics. ECE used 10 equal-width probability bins.")
    sg=subgroup[(subgroup.k.round(2)==0.10)&(subgroup.group=="income_group_traincut")]; low=sg[sg.value=="low"].iloc[0]; high=sg[sg.value=="high"].iloc[0]
    subgroup_text=(f"At 10% capacity, low-income respondents had a selection rate of {low.selection_rate:.3f} and recall of {low.recall_at_k:.3f}; high-income respondents had a selection rate of {high.selection_rate:.3f} and recall of {high.recall_at_k:.3f}. These values describe how one global queue distributed selection and missed cases; they do not identify the source of the differences or establish legal compliance.")
    tests={r.term:r for _,r in rsa[rsa.component=="surface_test"].iterrows()}
    rsa_text=(f"The response surface used a common 2.5-day center. The incongruence-line slope was {tests['incongruence_line_slope'].estimate:.4f} (95% CI: {tests['incongruence_line_slope'].ci_low:.4f}–{tests['incongruence_line_slope'].ci_high:.4f}), and its curvature was {tests['incongruence_line_curvature'].estimate:.4f} (95% CI: {tests['incongruence_line_curvature'].ci_low:.4f}–{tests['incongruence_line_curvature'].ci_high:.4f}). "
              f"The congruence-line curvature was negative ({tests['congruence_line_curvature'].estimate:.4f}, 95% CI: {tests['congruence_line_curvature'].ci_low:.4f}–{tests['congruence_line_curvature'].ci_high:.4f}), indicating that dissatisfaction also varied nonlinearly across matched WFH levels.")
    return {"[[ABSTRACT_RANKING_RESULTS]]":abstract_rank,"[[SAMPLE_RESULTS]]":sample_text,"[[ASSOCIATION_RESULTS]]":association,"[[MODEL_RESULTS]]":model,"[[CALIBRATION_RESULTS]]":calibration,"[[SUBGROUP_RESULTS]]":subgroup_text,"[[RSA_RESULTS]]":rsa_text}


def set_properties(doc, title):
    p=doc.core_properties
    p.title=title; p.subject="Directional hybrid-work misfit and job dissatisfaction"; p.author="Zidong Zhou; Rudzi Munap; Yue Meng"
    p.last_modified_by=p.author; p.keywords="hybrid work; work from home; social sustainability; job dissatisfaction; preference fit"
    p.comments=""; p.category="Article"


def build_manuscript():
    text=SOURCE.read_text(encoding="utf-8")
    for marker,replacement in rendered_blocks().items(): text=text.replace(marker,replacement)
    doc=new_mdpi_document(); set_properties(doc,"Directional Hybrid-Work Misfit and Job Dissatisfaction")
    sample=pd.read_csv(TABLES/"revision_sample_construction_flow.csv")
    desc=pd.read_csv(TABLES/"revision_table1_descriptives_by_misfit.csv")
    assoc=pd.read_csv(TABLES/"revision_association_model_grid.csv")
    perf=pd.read_csv(TABLES/"revision_ablation_temporal_performance.csv")
    rules=pd.read_csv(TABLES/"revision_strong_rule_baselines.csv")
    inserted=set(); front=True; current_section=""; skip_caption=False
    lines=text.splitlines(); i=0
    while i<len(lines):
        s=lines[i].strip(); i+=1
        if not s: continue
        if s.startswith("[["):
            if s=="[[TABLE_1]]" and s not in inserted: add_table_1(doc,sample)
            elif s=="[[TABLE_2]]" and s not in inserted: add_table_2(doc,desc)
            elif s=="[[TABLE_3]]" and s not in inserted: add_table_3(doc,assoc)
            elif s=="[[TABLE_4]]" and s not in inserted: add_table_4(doc,perf,rules)
            elif s=="[[TABLE_5]]" and s not in inserted: add_table(doc,"Table 5. Directional rule baselines.",rules)
            elif s=="[[FIGURE_1]]" and s not in inserted: add_figure(doc,"Figure 1. Directional strong misfit over time. Weighted prevalence is calculated among respondents with nonmissing desired and employer-planned WFH measures. Labels report the valid annual denominator; 2020 covers May–December and 2026 covers January–May. Strong misfit is a gap of at least two days.","fig1_directional_misfit_trend.png", width=4.5)
            elif s=="[[FIGURE_2]]" and s not in inserted: add_figure(doc,"Figure 2. Weighted job-dissatisfaction prevalence by planned-reference misfit group from July 2025 through February 2026.","fig2_job_dissatisfaction_by_misfit.png")
            elif s=="[[FIGURE_3]]" and s not in inserted: add_figure(doc,"Figure 3. Fixed-capacity ranking performance in the January–February 2026 temporal holdout. Bars show precision@10%, and labels show lift. The dashed line indicates the holdout outcome prevalence. The absolute-gap rule ranks the absolute desired–planned gap; the directional rule ranks the positive desired–planned gap; the additive rule combines the under-remote gap with prespecified commute, children, and income indicators. Conventional RF uses conventional characteristics; WFH RF adds desired and employer-planned WFH; Reduced RF excludes gender and children. The broader planned-only sensitivity model is reported in Supplementary Table S13.","fig3_capacity_constrained_performance.png")
            elif s=="[[FIGURE_4]]" and s not in inserted: add_figure(doc,"Figure 4. Reliability curves for the reduced-feature random forest. Points compare mean predicted probability with observed prevalence in 10 equal-width bins; the diagonal indicates perfect calibration. Brier scores are shown in the panel.","fig4_calibration.png")
            elif s=="[[FIGURE_5]]" and s not in inserted: add_figure(doc,"Figure 5. Subgroup selection and recall at 10% capacity. Bars show selection rates; points and intervals show recall. The dashed line marks overall recall. The gender labels reproduce the binary source coding (female = 1 and female = 0); education labels follow the SWAA categories shown in the figure. Income groups were defined using tertile cut points estimated in the 2025 training sample.","fig5_subgroup_audit.png")
            inserted.add(s); skip_caption=True; continue
        if skip_caption and (re.match(r"^(Figure|Table)\s+\d+\.",s) or s.startswith("Note:")):
            continue
        skip_caption=False
        if s.startswith("# "):
            add_p(doc,"Article",STYLE["article"])
            add_p(doc,s[2:],STYLE["title"])
            continue
        if s=="Type of the Paper: Article": continue
        if front and re.match(r"^Zidong Zhou",s): add_p(doc,s,STYLE["authors"]); continue
        if front and (re.match(r"^[123] ",s) or s.startswith("Correspondence:")): add_p(doc,s,STYLE["affiliation"]); continue
        if s=="## Abstract":
            # consume next nonblank paragraph and keywords
            while i<len(lines) and not lines[i].strip(): i+=1
            abstract=lines[i].strip(); i+=1
            add_labeled_p(doc,"Abstract: ",abstract,STYLE["abstract"])
            while i<len(lines) and not lines[i].strip(): i+=1
            if i<len(lines) and lines[i].strip().startswith("Keywords:"):
                kw=lines[i].strip()[len("Keywords:"):].strip(); i+=1
                add_labeled_p(doc,"Keywords: ",kw,STYLE["keywords"])
            front=False; continue
        if s.startswith("## "):
            current_section=s[3:]
            add_p(doc,current_section,STYLE["h1"],keep_with_next=True); continue
        if s.startswith("### "):
            add_p(doc,s[4:],STYLE["h2"],keep_with_next=True); continue
        if re.match(r"^RQ[1-4]\.",s):
            m=re.match(r"^(RQ[1-4]\.)\s*(.*)$",s)
            p=doc.add_paragraph(style=STYLE["text_no_indent"] if has_style(doc,STYLE["text_no_indent"]) else None)
            r=p.add_run(m.group(1)+" "); r.bold=True; p.add_run(m.group(2)); continue
        if current_section=="References" and re.match(r"^\d+\.",s):
            add_p(doc, re.sub(r"^\d+\.\s*", "", s), STYLE["reference"]); continue
        back_labels=["Supplementary Materials:","Author Contributions:","Funding:","Acknowledgments:","Institutional Review Board Statement:","Informed Consent Statement:","Data Availability Statement:","Code Availability Statement:","Conflicts of Interest:"]
        matched=next((x for x in back_labels if s.startswith(x)),None)
        if matched:
            add_labeled_p(doc,matched+" ",s[len(matched):].strip(),STYLE["back"]); continue
        if current_section=="Abbreviations": add_p(doc,s,STYLE["notes"]); continue
        add_p(doc,s,STYLE["text"])
    out=MAN/"Manuscript_MDPI_Sustainability_FINAL.docx"; MAN.mkdir(parents=True,exist_ok=True); doc.save(out); return out


def build_supplement():
    doc=new_mdpi_document(); set_properties(doc,"Supplementary Material")
    add_p(doc,"Supplementary Material",STYLE["article"])
    add_p(doc,"Directional Hybrid-Work Misfit and Job Dissatisfaction: Implications for Socially Sustainable Hybrid-Work Design",STYLE["title"])
    add_p(doc,"Zidong Zhou, Rudzi Munap and Yue Meng",STYLE["authors"])
    add_p(doc,"This file contains Supplementary Figures S1–S2 and an index to the complete spreadsheet tables. Tables S1–S43 and Additional Tables A1–A3 are provided in Supplementary File S2. Detailed variable definitions, model specifications, and robustness results are reported in those tables. Source CSV files are provided in Supplementary File S3.",STYLE["text_no_indent"])
    add_figure(doc,"Figure S1. Job dissatisfaction across signed planned-reference misfit bins. Positive values indicate that desired WFH exceeds employer-planned WFH.","figS1_misfit_bin_prevalence.png")
    add_figure(doc,"Figure S2. Planned-only response surface. Desired and employer-planned WFH are centered at the common 2.5-day midpoint. Predictions are masked where the nearest quarter-day cell contains fewer than 30 observations; point size represents local support.","figS2_planned_response_surface.png")
    idx=pd.DataFrame([(a,b) for a,b,_ in SUPPLEMENTARY_FILES],columns=["Supplementary item","CSV source"])
    add_table(doc,"Supplementary table index.",idx,"The spreadsheet workbook contains one formatted worksheet per listed table.")
    out=MAN/"Supplementary_Material_FINAL.docx"; MAN.mkdir(parents=True,exist_ok=True); doc.save(out); return out


def build_cover_letter():
    doc=Document(); sec=doc.sections[0]; sec.top_margin=sec.bottom_margin=Inches(.75); sec.left_margin=sec.right_margin=Inches(.9)
    normal=doc.styles["Normal"]; normal.font.name="Times New Roman"; normal.font.size=Pt(11); normal.paragraph_format.space_after=Pt(7)
    set_properties(doc,"Cover Letter")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Cover Letter"); r.bold=True; r.font.size=Pt(15)
    paragraphs=[
        "Dear Editors of Sustainability,",
        'We submit the manuscript “Directional Hybrid-Work Misfit and Job Dissatisfaction: Implications for Socially Sustainable Hybrid-Work Design” for consideration as an Article in Sustainability.',
        "The study examines whether the direction of the gap between desired and employer-planned work from home is associated with job dissatisfaction. The fully adjusted association model used 15,355 complete cases from September–December 2025. In that model, strong under-remote misfit was associated with an 11.5-percentage-point higher probability of job dissatisfaction. The adjusted under–over contrast was 7.9 percentage points (95% CI: 3.1–12.7).",
        "Adding desired and planned WFH increased random-forest precision@10% from 0.1706 to 0.2408; the paired-bootstrap increase was 7.02 percentage points (95% interval: 4.21–9.83). At 10% capacity, the under-remote rule achieved a precision of 0.2322, compared with 0.2505 for the reduced-feature random forest. The 1.84-percentage-point difference had a paired-bootstrap interval of −0.32 to 3.78 percentage points, so the manuscript does not claim equivalence or non-inferiority.",
        "Existing studies have mainly compared preferred with realized telework under pandemic-period conditions. The present manuscript extends this evidence by examining desired versus employer-planned WFH in a later U.S. repeated cross-section, formally comparing the two directions of misfit, and evaluating transparent temporal-holdout ranking. The manuscript fits the scope of Sustainability because it examines how employee voice, wellbeing, and accountable decision support can inform socially sustainable hybrid-work design.",
        "We reran the full statistical and document-generation pipeline from the verified SWAA archive in a clean directory, synchronized the manuscript with all generated tables and figures, and added automated checks for sample timing, model identity, reference definitions, calibration, and response-surface centering.",
        "We confirm that this manuscript and all parts of its content are neither published elsewhere nor under consideration by another journal. All authors have approved the manuscript and agree with its submission to Sustainability. The authors declare no conflicts of interest. Raw SWAA data are not redistributed; the code package records the data release, file hash, and reproduction command.",
        "Sincerely,", "Zidong Zhou", "Rudzi Munap", "Yue Meng",
    ]
    for t in paragraphs: doc.add_paragraph(t)
    out=MAN/"Cover_Letter_MDPI_Sustainability_FINAL.docx"; doc.save(out); return out


def main():
    parser=argparse.ArgumentParser(); parser.parse_args()
    for f in [build_manuscript(),build_supplement(),build_cover_letter()]: print(f.relative_to(ROOT))

if __name__=="__main__": main()
