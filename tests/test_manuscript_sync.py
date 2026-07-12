from __future__ import annotations

import os
import unittest
from pathlib import Path
from docx import Document


ROOT = Path(os.environ.get("SWAA_OUTPUTS_ROOT", Path(__file__).resolve().parents[1])).resolve()
MANUSCRIPT_SOURCE = ROOT / "documentation" / "current_sources" / "Manuscript_MDPI_Sustainability_source.md"
DOC_GENERATOR = ROOT / "code" / "generate_submission_documents.py"


class ManuscriptSyncTests(unittest.TestCase):
    def test_manuscript_references_new_supplement_range(self):
        text = MANUSCRIPT_SOURCE.read_text(encoding="utf-8")
        self.assertIn("Tables S1–S43", text)
        self.assertNotIn("Tables S1-S37", text)
        self.assertNotIn("S1–S37", text)

    def test_manuscript_uses_current_preference_and_model_terms(self):
        text = MANUSCRIPT_SOURCE.read_text(encoding="utf-8")
        for banned in [
            "preference-based voice",
            "minimal voice",
            "weak voice",
            "voice variables",
            "voice signals",
            "anti-leakage",
            "This asymmetry matters for Sustainability",
        ]:
            self.assertNotIn(banned, text)
        self.assertIn("desired and employer-planned WFH", text)
        self.assertIn("reduced-feature RF", text)
        self.assertIn("conventional worker and job characteristics", text)
        for stale in ["HRIS-like", "safer random forest", "reduced-sensitive-feature", "reduced-proximity", "0.256", "0.024", "confirms that the machine-learning gain"]:
            self.assertNotIn(stale, text)
        self.assertIn("paired-bootstrap interval", text)
        self.assertIn("does not establish equivalence", text)
        self.assertIn("same 2.5-day midpoint", text)


    def test_submission_copy_has_no_internal_placeholders(self):
        text = MANUSCRIPT_SOURCE.read_text(encoding="utf-8")
        for banned in ["AUTHOR ACTION REQUIRED", "Do not submit this bracketed text", "captures most of the observed"]:
            self.assertNotIn(banned, text)
        self.assertNotIn("comparable performance", text)
        self.assertNotIn("information-compression tools", text)
        self.assertNotIn("H1.", text)
        self.assertNotIn("H2.", text)
        for rq in range(1, 5):
            self.assertIn(f"RQ{rq}.", text)
        self.assertIn("poorly calibrated", text)
        self.assertIn("Institutional Review Board Statement: Ethical review and approval were waived", text)
        self.assertIn("Additional Tables A1–A3", text)


    def test_primary_association_time_window_is_explicit(self):
        text = MANUSCRIPT_SOURCE.read_text(encoding="utf-8")
        self.assertIn("September–December 2025", text)
        self.assertIn("15,355 of 37,434", text)
        self.assertIn("41.0%", text)
        self.assertNotIn("The primary association analysis used 2025–2026 data", text)
        self.assertIn("JVICT-DA-EXEMPT-2026-004", text)

    def test_cover_letter_generator_uses_current_terms(self):
        text = DOC_GENERATOR.read_text(encoding="utf-8")
        self.assertNotIn("preference-based voice", text)
        self.assertNotIn("anti-leakage", text)
        self.assertIn("does not claim equivalence or non-inferiority", text)
        self.assertIn("reduced-feature random forest", text)

    def test_submission_requirements_and_data_citation_are_explicit(self):
        text = MANUSCRIPT_SOURCE.read_text(encoding="utf-8")
        self.assertIn("Python 3.13.5 using pandas 2.2.3", text)
        self.assertIn("available to registered researchers", text)
        self.assertIn("Why Working from Home Will Stick", text)
        self.assertIn("Tables S1–S43 and Additional Tables A1–A3, whose complete titles", text)
        self.assertIn("Income groups were defined using tertile cut points estimated in the 2025 training sample", text)
        self.assertNotIn("publicly available, de-identified third-party survey data", text)

    def test_reference_numbering_is_continuous_and_ordered(self):
        text = MANUSCRIPT_SOURCE.read_text(encoding="utf-8")
        refs = text.split("## References", 1)[1]
        nums = [int(line.split(".",1)[0]) for line in refs.splitlines() if line and line.split(".",1)[0].isdigit()]
        self.assertEqual(nums, list(range(1, 49)))
        self.assertIn("[41,42]", text)
        self.assertIn("[47,48]", text)

    def test_generated_supplement_uses_full_title(self):
        path = ROOT / "manuscript" / "Supplementary_Material_FINAL.docx"
        if not path.exists():
            self.skipTest("generated supplementary material is not present; run code/run_all.py first")
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Directional Hybrid-Work Misfit and Job Dissatisfaction: Implications for Socially Sustainable Hybrid-Work Design", text)

    def test_generated_manuscript_contains_all_core_tables_and_figures(self):
        path = ROOT / "manuscript" / "Manuscript_MDPI_Sustainability_FINAL.docx"
        if not path.exists():
            self.skipTest("generated manuscript is not present; run code/run_all.py first")
        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)
        for n in range(1, 6):
            self.assertIn(f"Figure {n}.", text)
            self.assertIn(f"Table {n}.", text)
        self.assertGreaterEqual(len(doc.inline_shapes), 5)


if __name__ == "__main__":
    unittest.main()
